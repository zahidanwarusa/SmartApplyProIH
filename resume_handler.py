import json
import logging
import time
import re
import os
from pathlib import Path
from typing import Dict, Optional, List, Union
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import docx.oxml.shared
from docx.opc.constants import RELATIONSHIP_TYPE

from config import DEFAULT_RESUME, RESUME_DIR
from gemini_service import GeminiService

class ResumeHandler:
    """Handles resume generation and optimization for the new v2 format"""
    
    def __init__(self):
        self.gemini = GeminiService()
        self.logger = logging.getLogger(__name__)
        
    def generate_resume(self, job_details: Dict) -> Optional[str]:
        """Generate optimized resume for a job using the new v2 format"""
        try:
            # Load default resume
            if not DEFAULT_RESUME.exists():
                self.logger.error("Default resume template not found")
                return None
                
            with open(DEFAULT_RESUME, 'r', encoding='utf-8') as f:  # FIX: Explicit UTF-8 encoding
                resume_data = json.load(f)
                
            # Create output directories
            RESUME_DIR.mkdir(parents=True, exist_ok=True)
            
            # Log details for debugging
            self.logger.info(f"Processing resume for job: {job_details.get('title', 'Unknown')}")
            
            # Process sections for new v2 format
            sections_to_process = [
                'professional_summary',
                'core_competencies', 
                'professional_experience'
            ]
            
            for section_name in sections_to_process:
                if section_name in resume_data:
                    self.logger.info(f"Optimizing {section_name}...")
                    try:
                        # CRITICAL FIX: Store original content before optimization
                        original_content = resume_data[section_name].copy() if isinstance(resume_data[section_name], dict) else resume_data[section_name][:]
                        
                        updated_section = self.gemini.optimize_resume_section(
                            section_name,
                            resume_data[section_name],
                            job_details
                        )
                        
                        if updated_section:
                            # Deep comparison with original before replacing
                            original_normalized = json.dumps(self._normalize_content(original_content), sort_keys=True)
                            updated_normalized = json.dumps(self._normalize_content(updated_section), sort_keys=True)
                            
                            if original_normalized != updated_normalized:
                                resume_data[section_name] = updated_section
                                self.logger.info(f"Successfully updated {section_name} with meaningful changes")
                            else:
                                self.logger.warning(f"No significant changes detected for {section_name}")
                                resume_data[section_name] = original_content
                        else:
                            self.logger.warning(f"No valid response for {section_name}, keeping original")
                            resume_data[section_name] = original_content
                            
                    except Exception as e:
                        self.logger.error(f"Error updating {section_name}: {str(e)}")
                        resume_data[section_name] = original_content
                    
                    # Avoid rate limiting
                    time.sleep(2)
            
            # VALIDATION: Ensure all required v2 sections are present
            required_sections = ['header', 'professional_summary', 'core_competencies', 'professional_experience', 'education']
            missing_sections = [section for section in required_sections if section not in resume_data]
            
            if missing_sections:
                self.logger.error(f"Missing required sections: {missing_sections}")
                return None
            
            # Generate resume filename
            base_filename = self._create_professional_filename(job_details)
            
            # Ensure filename is unique
            resume_filename = self._ensure_unique_filename(base_filename, ".docx")
            json_filename = self._ensure_unique_filename(base_filename, ".json")
            
            # Generate files
            resume_path = RESUME_DIR / resume_filename
            json_path = RESUME_DIR / json_filename
            
            # Save JSON for reference with UTF-8 encoding
            with open(json_path, 'w', encoding='utf-8') as f:  # FIX: Explicit UTF-8 encoding
                json.dump(resume_data, f, indent=2, ensure_ascii=False)
                
            # Convert to DOCX using updated ResumeConverter
            converter = ResumeConverter()
            converter.convert_resume(resume_data)
            converter.save(str(resume_path))
            
            self.logger.info(f"Resume saved to {resume_path}")
            return str(resume_path)
            
        except Exception as e:
            self.logger.error(f"Error generating resume: {str(e)}")
            return None
            
    def _normalize_content(self, content):
        """Normalize content for comparison by removing formatting markers"""
        if isinstance(content, list):
            return [self._normalize_text(item) if isinstance(item, str) else self._normalize_content(item) for item in content]
        elif isinstance(content, dict):
            return {key: self._normalize_content(value) for key, value in content.items()}
        elif isinstance(content, str):
            return self._normalize_text(content)
        else:
            return content
    
    def _normalize_text(self, text):
        """Normalize text by removing bold markers and extra whitespace"""
        if not isinstance(text, str):
            return text
        # Remove bold markers
        text = re.sub(r'\*\*', '', text)
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
            
    def _create_professional_filename(self, job_details: Dict) -> str:
        """Create a professional resume filename based on job details"""
        job_title = job_details.get('title', 'Job').strip()
        role_keywords = ['SDET', 'QA', 'Test', 'Quality', 'Automation', 'Engineer', 'Lead', 'Senior', 'Manager', 'Business', 'Analyst', 'BA', 'Tester']
        
        role_parts = []
        for keyword in role_keywords:
            if re.search(r'\b' + re.escape(keyword) + r'\b', job_title, re.IGNORECASE):
                role_parts.append(keyword)
        
        if not role_parts:
            role_type = "QA"
        else:
            role_type = "_".join(role_parts)
        
        return f"{role_type}_Resume_Zahid_Anwar"
        
    def _ensure_unique_filename(self, base_filename: str, extension: str) -> str:
        """Ensure filename is unique by adding a counter if needed"""
        filename = f"{base_filename}{extension}"
        counter = 1
        
        while (RESUME_DIR / filename).exists():
            filename = f"{base_filename}_v{counter}{extension}"
            counter += 1
            
        return filename


class ResumeConverter:
    """FIXED: Converts resume JSON to formatted DOCX document using new v2 format"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
        
    def _setup_document(self):
        """Setup document properties and styles with correct fonts"""
        # Set up margins (1 inch all around)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Customize heading styles
        styles = self.doc.styles
        
        # Heading 1 style (for main sections)
        heading1_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Tahoma'  # As requested
        heading1_font.size = Pt(14)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(0, 51, 102)  # Professional blue
        
        # Normal style (for body text)
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'  # As requested
        normal_font.size = Pt(11)

    def _create_header(self, header_data: Dict):
        """Create professional header section"""
        # Name (larger, bold)
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(header_data['name'])
        name_run.font.name = 'Tahoma'
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Contact info (centered, smaller)
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_info = f"{header_data['email']} | {header_data['phone']} | {header_data.get('linkedin', '')} | {header_data['citizenship']}"
        contact_run = contact_para.add_run(contact_info)
        contact_run.font.name = 'Calibri'
        contact_run.font.size = Pt(11)
        
        # Add spacing after header
        contact_para.paragraph_format.space_after = Pt(18)

    def _add_professional_summary(self, summary_data: Dict):
        """Add unified professional summary section (v2 format)"""
        heading = self.doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        heading.style = self.doc.styles['Heading 1']
        for run in heading.runs:
            run.font.name = 'Tahoma'
            run.font.bold = True
        
        # Create formatted paragraphs for each component
        for field in ['title_experience', 'track_record', 'expertise', 'core_value']:
            if field in summary_data and summary_data[field]:
                p = self.doc.add_paragraph()
                self._add_formatted_text(p, summary_data[field])
                p.paragraph_format.space_after = Pt(8)
        

    def _add_core_competencies(self, competencies_data: Dict):
        """FIXED: Add core competencies section with bullet-separated format"""
        heading = self.doc.add_heading('CORE COMPETENCIES', level=1)
        heading.style = self.doc.styles['Heading 1']
        for run in heading.runs:
            run.font.name = 'Tahoma'
            run.font.bold = True
        
        # FIXED: Clean category display names
        category_display_names = {
            'programming_and_automation': 'Programming & Automation',
            'testing_frameworks': 'Testing Frameworks',
            'cloud_and_devops': 'Cloud & DevOps',
            'api_and_performance': 'API & Performance',
            'quality_tools': 'Quality Tools',
            'databases': 'Databases',
            'domain_expertise': 'Domain Expertise',
            'leadership': 'Leadership'
        }
        
        # Format each competency category
        for category, skills in competencies_data.items():
            if skills:  # Only add if there are skills
                p = self.doc.add_paragraph()
                
                # Category name (bold, clean formatting)
                display_name = category_display_names.get(category, category.replace('_', ' ').title())
                category_run = p.add_run(f"{display_name}: ")
                category_run.font.name = 'Calibri'
                category_run.font.bold = True
                
                # FIXED: Skills with proper bullet separation and bold handling
                for i, skill in enumerate(skills):
                    if i > 0:
                        # FIXED: Use proper bullet character (•) with UTF-8 support
                        bullet_run = p.add_run(' • ')
                        bullet_run.font.name = 'Calibri'
                    
                    # FIXED: Handle bold formatting correctly
                    self._add_formatted_text(p, skill)
                
                p.paragraph_format.space_after = Pt(6)
        

    def _add_professional_experience(self, experience_data: List[Dict]):
        """FIXED: Add professional experience section with v2 format"""
        heading = self.doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        heading.style = self.doc.styles['Heading 1']
        for run in heading.runs:
            run.font.name = 'Tahoma'
            run.font.bold = True
        
        for exp in experience_data:
            # Company and Location (bold)
            company_para = self.doc.add_paragraph()
            company_run = company_para.add_run(f"{exp['company']} | {exp['location']}")
            company_run.font.name = 'Calibri'
            company_run.font.bold = True
            company_run.font.size = Pt(13)
            
            # Position and Duration
            position_para = self.doc.add_paragraph()
            position_run = position_para.add_run(f"{exp['position']} | {exp['duration']}")
            position_run.font.name = 'Calibri'
            position_run.font.bold = True
            
            # Summary
            if 'summary' in exp and exp['summary']:
                summary_para = self.doc.add_paragraph()
                self._add_formatted_text(summary_para, exp['summary'])
                summary_para.paragraph_format.space_after = Pt(8)
            
            # Key achievements (if present)
            if 'key_achievements' in exp:
                for achievement in exp['key_achievements']:
                    bullet_para = self.doc.add_paragraph()
                    bullet_para.style = 'List Bullet'
                    self._add_formatted_text(bullet_para, achievement)
                    bullet_para.paragraph_format.space_after = Pt(4)
            
            # Detailed achievements (if present)
            if 'detailed_achievements' in exp:
                for achievement in exp['detailed_achievements']:
                    bullet_para = self.doc.add_paragraph()
                    bullet_para.style = 'List Bullet'
                    self._add_formatted_text(bullet_para, achievement)
                    bullet_para.paragraph_format.space_after = Pt(4)
            
            # FIXED: Environment/Skills section - clean technical tools only
            if 'environment' in exp and exp['environment']:
                env_para = self.doc.add_paragraph()
                env_run = env_para.add_run('Environment: ')
                env_run.font.name = 'Calibri'
                env_run.font.bold = True
                
                # FIXED: Clean environment text - remove narrative phrases
                clean_environment = self._clean_environment_text(exp['environment'])
                self._add_formatted_text(env_para, clean_environment)
                    

    def _clean_environment_text(self, environment_text: str) -> str:
        """FIXED: Clean environment text to show only technical tools"""
        if not environment_text:
            return environment_text
            
        # Remove narrative phrases that don't belong in environment section
        narrative_phrases = [
            r'demonstrating \*\*[^*]+\*\* and [^.]+\.',
            r'showcasing \*\*[^*]+\*\* and [^.]+\.',
            r'contributing to \*\*[^*]+\*\*[^.]*\.',
            r'reflecting \*\*[^*]+\*\* and [^.]+\.',
            r'showing \*\*[^*]+\*\* and [^.]+\.',
            r', demonstrating[^.]+\.',
            r', showcasing[^.]+\.',
            r', contributing[^.]+\.',
            r', reflecting[^.]+\.',
        ]
        
        cleaned = environment_text
        for pattern in narrative_phrases:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
        
        # Clean up any remaining artifacts
        cleaned = re.sub(r',\s*$', '', cleaned)  # Remove trailing commas
        cleaned = re.sub(r'\s+', ' ', cleaned)   # Normalize whitespace
        
        return cleaned.strip()

    def _add_education(self, education_data: List[Dict]):
        """Add education section"""
        heading = self.doc.add_heading('EDUCATION', level=1)
        heading.style = self.doc.styles['Heading 1']
        for run in heading.runs:
            run.font.name = 'Tahoma'
            run.font.bold = True
        
        for edu in education_data:
            p = self.doc.add_paragraph()
            
            # Format: MBA, Information Technology | Strayer University, USA (2016)
            degree_run = p.add_run(f"{edu['degree']}, {edu['major']}")
            degree_run.font.name = 'Calibri'
            degree_run.font.bold = True
            
            separator_run = p.add_run(' | ')
            separator_run.font.name = 'Calibri'
            
            university_run = p.add_run(f"{edu['university']} ({edu['year']})")
            university_run.font.name = 'Calibri'
            
            p.paragraph_format.space_after = Pt(6)

    def _add_formatted_text(self, paragraph, text):
        """FIXED: Add text to paragraph with proper bold formatting"""
        if not text:
            return
        
        # FIXED: Handle bold highlighting with ** markers
        # Split by ** but keep empty strings to maintain positioning
        parts = text.split('**')
        
        for i, part in enumerate(parts):
            if part or i == 0:  # Include first part even if empty
                run = paragraph.add_run(part)
                # FIXED: Odd indices (1, 3, 5...) should be bold
                run.bold = i % 2 == 1
                run.font.name = 'Calibri'

    def convert_resume(self, resume_data: Dict):
        """Convert full resume from JSON data to DOCX using new v2 format"""
        # Create header
        self._create_header(resume_data['header'])
        
        # Add Professional Summary (new unified section)
        self._add_professional_summary(resume_data['professional_summary'])
        
        # Add Core Competencies (replaces technical skills)
        self._add_core_competencies(resume_data['core_competencies'])
        
        # Add Professional Experience
        self._add_professional_experience(resume_data['professional_experience'])
        
        # Add Education
        self._add_education(resume_data['education'])

    def save(self, filename: str):
        """FIXED: Save the document with UTF-8 support"""
        try:
            self.doc.save(filename)
        except Exception as e:
            # If there are encoding issues, try to fix them
            logging.error(f"Error saving document: {e}")
            raise
